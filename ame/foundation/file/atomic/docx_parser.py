"""
DOCX 文件解析器

支持:
- .docx 文件
- 标题、段落识别
- 表格提取
- 列表提取

依赖: python-docx
"""

import logging
from typing import List
from pathlib import Path

from .base import FileParserBase
from ..core.models import ParsedDocument, DocumentSection, DocumentFormat, SectionType
from ..core.exceptions import DependencyMissingError, ParseError

logger = logging.getLogger(__name__)


class DocxParser(FileParserBase):
    """
    DOCX 解析器
    
    特性:
    - 提取标题、段落
    - 提取表格
    - 提取列表
    - 保留样式信息
    
    依赖:
    - python-docx
    """
    
    SUPPORTED_EXTENSIONS = {"docx"}
    
    def __init__(self):
        """初始化 DOCX 解析器"""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """检查依赖是否安装"""
        try:
            import docx
            self.docx_lib = docx
        except ImportError:
            logger.error("DOCX 解析依赖未安装")
            self.docx_lib = None
    
    def can_parse(self, file_path: str) -> bool:
        """判断是否支持该文件"""
        extension = self._get_file_extension(file_path)
        return extension in self.SUPPORTED_EXTENSIONS and self.docx_lib is not None
    
    async def parse(self, file_path: str) -> ParsedDocument:
        """
        解析 DOCX 文件
        
        Args:
            file_path: 文件路径
        
        Returns:
            parsed_doc: 解析结果
        
        Raises:
            ImportError: python-docx 未安装
        """
        if self.docx_lib is None:
            raise DependencyMissingError(
                "python-docx",
                "pip install python-docx"
            )
        
        # 验证文件
        path = self._validate_file_exists(file_path)
        
        try:
            return await self._parse_docx(path)
        
        except Exception as e:
            logger.error(f"DOCX 文件解析失败: {file_path}, {e}", exc_info=True)
            raise
    
    async def _parse_docx(self, path: Path) -> ParsedDocument:
        """
        解析 DOCX 文档
        
        Args:
            path: 文件路径
        
        Returns:
            parsed_doc: 解析结果
        """
        import docx
        
        doc = docx.Document(path)
        
        sections = []
        raw_content_parts = []
        
        # 遍历文档元素
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            raw_content_parts.append(text)
            
            # 判断是否是标题
            if para.style.name.startswith("Heading"):
                level = self._extract_heading_level(para.style.name)
                
                sections.append(DocumentSection(
                    type=self._get_heading_type(level),
                    content=text,
                    level=level,
                    metadata={
                        "style": para.style.name,
                        "bold": para.runs[0].bold if para.runs else False,
                    }
                ))
            
            # 判断是否是列表
            elif self._is_list_paragraph(para):
                sections.append(DocumentSection(
                    type=SectionType.LIST_ITEM,
                    content=text,
                    metadata={"style": para.style.name}
                ))
            
            # 普通段落
            else:
                sections.append(DocumentSection(
                    type=SectionType.PARAGRAPH,
                    content=text,
                    metadata={"style": para.style.name}
                ))
        
        # 提取表格
        for table_idx, table in enumerate(doc.tables):
            table_text = self._extract_table(table)
            
            sections.append(DocumentSection(
                type=SectionType.TABLE,
                content=table_text,
                metadata={
                    "table_index": table_idx,
                    "rows": len(table.rows),
                    "cols": len(table.columns)
                }
            ))
            
            raw_content_parts.append(table_text)
        
        raw_content = "\n\n".join(raw_content_parts)
        
        # 提取文档属性
        core_props = doc.core_properties
        
        return ParsedDocument(
            format=DocumentFormat.DOCX,
            file_path=str(path),
            raw_content=raw_content,
            sections=sections,
            metadata={
                "parser": "python-docx",
                "file_size": path.stat().st_size,
                "author": core_props.author if core_props.author else None,
                "title": core_props.title if core_props.title else None,
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
            }
        )
    
    def _extract_heading_level(self, style_name: str) -> int:
        """
        从样式名提取标题级别
        
        Args:
            style_name: 样式名（如 "Heading 1"）
        
        Returns:
            level: 标题级别（1-6）
        """
        try:
            # 尝试从 "Heading 1" 中提取数字
            if "Heading" in style_name:
                parts = style_name.split()
                if len(parts) >= 2 and parts[1].isdigit():
                    level = int(parts[1])
                    return min(level, 6)  # 最多支持 6 级标题
        except Exception:
            pass
        
        return 1  # 默认为 1 级标题
    
    def _get_heading_type(self, level: int) -> SectionType:
        """将标题级别转换为 SectionType"""
        mapping = {
            1: SectionType.HEADING_1,
            2: SectionType.HEADING_2,
            3: SectionType.HEADING_3,
            4: SectionType.HEADING_4,
            5: SectionType.HEADING_5,
            6: SectionType.HEADING_6,
        }
        return mapping.get(level, SectionType.HEADING_1)
    
    def _is_list_paragraph(self, para) -> bool:
        """
        判断段落是否是列表项
        
        Args:
            para: 段落对象
        
        Returns:
            is_list: 是否是列表
        """
        # 检查样式名
        if "List" in para.style.name:
            return True
        
        # 检查编号格式（有些列表可能没有 List 样式）
        if para.text.strip().startswith(("•", "-", "*", "○", "■")):
            return True
        
        return False
    
    def _extract_table(self, table) -> str:
        """
        提取表格内容
        
        Args:
            table: 表格对象
        
        Returns:
            table_text: 表格文本
        """
        lines = []
        
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
        
        return "\n".join(lines)
